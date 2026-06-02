from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colores como tuplas (r, g, b) ─────────────────────────────
VERDE   = (0x1a, 0x6b, 0x3a)
DORADO  = (0xc9, 0xa2, 0x27)
CAFE    = (0x5c, 0x33, 0x17)
AZUL    = (0x1a, 0x3a, 0x6b)
ROJO    = (0xb9, 0x1c, 0x1c)
GRIS    = (0x55, 0x55, 0x55)
NEGRO   = (0x1a, 0x1a, 0x1a)
BLANCO  = (0xff, 0xff, 0xff)

def rgb(c):
    return RGBColor(c[0], c[1], c[2])

def hex6(c):
    return "{:02X}{:02X}{:02X}".format(c[0], c[1], c[2])

PERSONAS = {
    1: ("Alejandro Aguilar",  VERDE),
    2: ("Alejandro Cabrera",  DORADO),
    3: ("Luis Pertuz",        CAFE),
    4: ("Jesus Ramirez",      AZUL),
}

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(3.0)


def font(run, size=11, bold=False, italic=False, color=NEGRO):
    run.font.name   = "Calibri"
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def heading(text, level=1, color=VERDE):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    sz  = {1: 22, 2: 16, 3: 13}.get(level, 12)
    font(run, sz, bold=True, color=color)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)


def divider(color=DORADO):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr    = p._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex6(color))
    pBdr.append(bottom)
    pPr.append(pBdr)


def cell_shade(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def speaker(num, text, datos=""):
    name, color = PERSONAS[num]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    tag = p.add_run(f"[{name.split()[0].upper()}]  ")
    font(tag, 10, bold=True, color=color)
    body = p.add_run(text)
    font(body, 11, color=NEGRO)
    if datos:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent  = Cm(1.8)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(2)
        d = p2.add_run(f"  {datos}")
        font(d, 9.5, italic=True, color=GRIS)


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"[NOTA]  {text}")
    font(r, 9, italic=True, color=(0x88, 0x88, 0x88))


def section_tag(num, title, color=VERDE):
    p = doc.add_paragraph()
    r1 = p.add_run(f"  SEC. {num:02d}  ")
    font(r1, 9, bold=True, color=BLANCO)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6(color))
    r1._r.get_or_add_rPr().append(shd)
    r2 = p.add_run(f"  {title.upper()}")
    font(r2, 14, bold=True, color=color)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)


# ════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\nANÁLISIS ECONÓMICO · RISARALDA")
font(r, 28, bold=True, color=VERDE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Guión de presentación — 4 expositores")
font(r, 14, color=DORADO)

divider(DORADO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Fundamentos de Economía · 2026\n")
font(r, 11, italic=True, color=GRIS)

for num, (name, color) in PERSONAS.items():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"Expositor {num}:  ")
    font(r1, 11, color=GRIS)
    r2 = p.add_run(name)
    font(r2, 11, bold=True, color=color)

divider(VERDE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Duración estimada: 12 – 15 minutos")
font(r, 10, italic=True, color=GRIS)

doc.add_page_break()

# ════════════════════════════════════════
#  INSTRUCCIONES
# ════════════════════════════════════════
heading("Instrucciones generales", 2, CAFE)
divider(CAFE)

for txt in [
    "El guión sigue las seis secciones de la página web en orden de aparición.",
    "Cada expositor tiene un color de etiqueta propio. Respetar los turnos de palabra.",
    "Los renglones con '  ' (indentados) son datos clave que se mencionan al señalar en pantalla.",
    "Los bloques '[NOTA]' son indicaciones de navegación o énfasis visual, no se leen en voz alta.",
    "Tiempos sugeridos: Apertura 1 min · Perfil 1,5 min · Social 2 min · Económico 3 min · Retos 2 min · Propuestas 2,5 min · Conclusiones 2 min.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(txt)
    font(r, 10.5, color=NEGRO)

doc.add_paragraph()

# ════════════════════════════════════════
#  APERTURA
# ════════════════════════════════════════
heading("APERTURA  ·  Hero de la página", 1, VERDE)
divider()

speaker(1,
    "Buenas tardes a todos. Hoy les presentamos el análisis económico del departamento de Risaralda, "
    "una región estratégica en el corazón del Eje Cafetero colombiano. Esta página web fue desarrollada "
    "como parte de la materia Fundamentos de Economía del 2026.")

speaker(2,
    "Risaralda ocupa apenas 4.140 km² — un territorio relativamente pequeño — pero concentra cerca "
    "de 997.000 habitantes, representa el 1,65% del PIB nacional y es uno de los departamentos con "
    "mejores indicadores sociales del país.",
    "Población ~997.000 (DANE 2025)  ·  PIB 1,65% nacional (DANE 2024)  ·  14 municipios")

speaker(3,
    "Su capital, Pereira, es el motor económico de la región y uno de los nodos logísticos más "
    "importantes de Colombia, conectando los ejes viales hacia el interior del país con los puertos "
    "del Pacífico.")

speaker(4,
    "Exploraremos seis secciones: perfil general, indicadores sociales, indicadores económicos, "
    "problemáticas, propuestas de los equipos asesores y las conclusiones del análisis. "
    "Los invitamos a seguir la exposición en la pantalla mientras navegamos la página.")

note("Hacer scroll suave hasta la sección «Perfil General».")

# ════════════════════════════════════════
#  SECCIÓN 01 — PERFIL GENERAL
# ════════════════════════════════════════
section_tag(1, "Perfil General", VERDE)
divider()

speaker(1,
    "Comenzamos con el perfil general. Risaralda fue creado en 1966 tras la separación del Viejo "
    "Caldas, y su identidad está profundamente ligada a la tradición cafetera. Su Paisaje Cultural "
    "Cafetero es Patrimonio de la Humanidad reconocido por la UNESCO.")

speaker(1,
    "Geográficamente, limita con Caldas al norte, Quindío al sur, Valle del Cauca al suroeste, "
    "Chocó al oeste y Antioquia al noroeste. Esta posición estratégica lo convierte en el corredor "
    "natural del Eje Cafetero.",
    "Extensión: 4.140 km²  ·  Región Andina  ·  Capital: Pereira")

speaker(2,
    "En demografía, la proyección del DANE 2025 estima cerca de 997.000 habitantes, con una "
    "densidad de 241 habitantes por kilómetro cuadrado. El 78% vive en zonas urbanas y el 22% "
    "en áreas rurales, según el Censo Nacional de Población y Vivienda de 2018.",
    "Población ~997.000 (DANE 2025)  ·  Densidad 241 hab/km²  ·  78% urbano / 22% rural (CNPV 2018)")

speaker(2,
    "Sus recursos naturales son excepcionales: suelos volcánicos de alta fertilidad, los ríos "
    "Otún y Cauca, el Parque Nacional Natural Los Nevados y una biodiversidad propia de los "
    "Bosques Andinos.")

note("Señalar los cuatro íconos de recursos naturales en pantalla.")

# ════════════════════════════════════════
#  SECCIÓN 02 — INDICADORES SOCIALES
# ════════════════════════════════════════
section_tag(2, "Indicadores Sociales", AZUL)
divider(AZUL)

speaker(3,
    "Pasamos a los indicadores sociales, donde Risaralda muestra un desempeño superior al "
    "promedio nacional en todas las dimensiones. Empecemos con el Índice de Desarrollo Humano.")

speaker(3,
    "El IDH de Risaralda es de 0,764, según el PNUD Colombia 2024, superior al promedio nacional "
    "de 0,754. Clasifica en la categoría de «alto desarrollo» y refleja avances simultáneos en "
    "salud, educación e ingresos.",
    "IDH Risaralda: 0,764  vs  Nacional: 0,754  (PNUD 2024)  ·  Categoría: Alto")

speaker(4,
    "La esperanza de vida es de 78,1 años en promedio — 75,3 para hombres y 80,9 para mujeres —, "
    "superior al promedio nacional de 77,5 años. La cobertura del SGSSS alcanza el 96,8% de la "
    "población, según datos de 2024.",
    "Esperanza de vida: 78,1 años (DANE 2024)  ·  Cobertura salud: 96,8% (SGSSS 2024)")

speaker(3,
    "En educación, la cobertura de básica llega al 93,1% y la tasa de analfabetismo es de solo "
    "3,4%, frente al 4,7% nacional. Datos del DANE 2024 confirman que Risaralda supera al "
    "promedio en todas las dimensiones educativas básicas.",
    "Cobertura básica: 93,1%  ·  Analfabetismo Risaralda: 3,4%  vs  Nacional: 4,7%  (DANE 2024)")

speaker(4,
    "Ahora los indicadores de pobreza. La pobreza monetaria en Risaralda bajó al 22,1% en 2024, "
    "mientras el promedio nacional es 32,7%. Una diferencia de 10,6 puntos porcentuales que "
    "demuestra el buen desempeño relativo del departamento.",
    "Pobreza monetaria: 22,1% Risaralda  vs  32,7% Nacional  (DANE 2024)  ·  Brecha: −10,6 pp")

speaker(3,
    "La pobreza multidimensional también es inferior: 8,2% departamental frente a 9,5% nacional. "
    "El coeficiente de Gini es 0,464 en Risaralda versus 0,541 nacional: menos desigual que el "
    "país, aunque todavía hay trabajo por hacer.",
    "Pobreza multidim.: 8,2%  vs  9,5%  ·  Gini: 0,464  vs  0,541  (DANE 2024)")

speaker(4,
    "Sobre servicios públicos, el sector urbano tiene coberturas casi universales: 98,5% en "
    "acueducto, 96,1% en alcantarillado y 99,3% en energía eléctrica. Sin embargo, el acueducto "
    "rural apenas cubre el 64% de la población, una brecha crítica que debe atenderse.",
    "Acueducto urbano: 98,5%  ·  Alcantarillado: 96,1%  ·  Energía: 99,3%  ·  Acueducto rural: 64%")

note("Señalar las barras de progreso y la advertencia del sector rural.")

# ════════════════════════════════════════
#  SECCIÓN 03 — INDICADORES ECONÓMICOS
# ════════════════════════════════════════
section_tag(3, "Indicadores Económicos", CAFE)
divider(CAFE)

speaker(1,
    "Entramos al corazón del análisis: los indicadores económicos. Risaralda aporta el 1,65% del "
    "PIB nacional, con un PIB per cápita de USD 6.380 según DANE 2024. Es un departamento de "
    "renta media dentro del contexto colombiano.",
    "PIB nacional: 1,65%  ·  PIB per cápita: USD 6.380  (DANE 2024)")

speaker(2,
    "El mercado laboral muestra mejoras recientes. La tasa de desempleo del área metropolitana "
    "de Pereira se ubica en 9,8% en 2025. Sin embargo, la informalidad laboral del 41,5% "
    "significa que casi la mitad de los trabajadores carece de seguridad social y pensión.",
    "Desempleo AM Pereira: 9,8% (DANE 2025)  ·  Informalidad: 41,5% (DANE 2024)")

speaker(1,
    "La gráfica de barras compara la pobreza monetaria: Risaralda 22,1% en verde versus el "
    "promedio nacional 32,7% en rojo. La gráfica de dona muestra la composición del PIB "
    "por sectores económicos.",
    "Composición PIB: Comercio/Hoteles/Transporte 21,2%  ·  Manufactura 12,1%  ·  Agricultura 10,5%  ·  Otros 56,2%")

speaker(2,
    "En comercio exterior, Risaralda exportó aproximadamente USD 342 millones en 2023, con "
    "un crecimiento acumulado del 54,6% entre 2019 y 2023. El café verde representa cerca del "
    "52% del total exportado, seguido por aguacate Hass, confecciones y ensamble automotriz.",
    "Exportaciones: USD ~342M (DANE 2023)  ·  Café ~52%  ·  Crecimiento 2019–2023: +54,6%")

speaker(1,
    "Los destinos principales en 2023 fueron Estados Unidos, Alemania, Japón, Bélgica y Países "
    "Bajos. El turismo aporta cerca del 4,5% del PIB con aproximadamente 1,2 millones de "
    "visitantes anuales según el Ministerio de Comercio.",
    "Destinos: EE.UU. · Alemania · Japón · Bélgica · Países Bajos  ·  Turismo: ~1,2M visitantes/año (MinCIT 2023)")

speaker(2,
    "Sobre la inflación: tras el pico de 2022 con 13,12%, hay una desinflación sostenida. "
    "En 2023 fue 9,28%, en 2024 bajó a 5,20%, y la proyección preliminar para 2025 es 4,32% "
    "nacional y 4,18% para Risaralda. La meta del Banco de la República es llegar a 3,0%.",
    "IPC: 2022→13,12%  ·  2023→9,28%  ·  2024→5,20%  ·  2025(p)→4,32%  ·  Meta BanRep: 3,0%")

note("Señalar la gráfica de líneas de inflación. La línea verde (Risaralda) siempre ligeramente por debajo de la roja (Colombia).")

# ════════════════════════════════════════
#  SECCIÓN 04 — PROBLEMÁTICAS
# ════════════════════════════════════════
section_tag(4, "Problemáticas y Retos", ROJO)
divider(ROJO)

speaker(3,
    "A pesar de los buenos indicadores, Risaralda enfrenta tres problemáticas estructurales "
    "que limitan su desarrollo y son el punto de partida para las propuestas que presentaremos.")

speaker(3,
    "Primera: el deterioro del mercado laboral urbano. Aunque el desempleo ha bajado al 9,8%, "
    "persiste una marcada brecha de género en la desocupación. La informalidad al 41,5% implica "
    "que casi la mitad de la fuerza laboral carece de seguridad social, pensión y estabilidad.",
    "Desempleo: 9,8%  ·  Informalidad: 41,5%  ·  Brecha de género significativa en desempleo")

speaker(4,
    "Segunda: la brecha urbano-rural. Mientras Pereira y Dosquebradas concentran inversión y "
    "servicios, municipios del occidente como Pueblo Rico, Mistrató y Guática enfrentan alta "
    "pobreza y bajo acceso a servicios. El 36% de la población rural sin acueducto es el "
    "indicador más crítico de esta desigualdad territorial.",
    "Acueducto rural: 64% cobertura  ·  Brecha ingresos y servicios urbano-rural")

speaker(3,
    "Tercera: la dependencia de precios agrícolas volátiles. El pequeño productor cafetero es "
    "vulnerable a las fluctuaciones del precio internacional del café y al encarecimiento de "
    "fertilizantes importados. Esta dependencia genera ciclos de auge y crisis que erosionan "
    "la estabilidad del ingreso campesino y frenan la inversión rural.")

note("Pausa breve. Estos tres retos son el puente directo hacia las propuestas de solución.")

# ════════════════════════════════════════
#  SECCIÓN 05 — PROPUESTAS
# ════════════════════════════════════════
section_tag(5, "Panel de Propuestas", VERDE)
divider()

speaker(4,
    "Pasamos al panel de propuestas. Cada equipo asesor plantea una estrategia de desarrollo "
    "regional que responde directamente a las problemáticas identificadas. Son cuatro propuestas "
    "complementarias.")

speaker(1,
    "Propuesta 1 — Innovación Tecnológica: crear un Hub de Desarrollo Tecnológico y Robótica "
    "en Pereira, financiado con recursos de regalías, que capacite a jóvenes en automatización "
    "industrial. El objetivo es tecnificar la producción, reducir el desempleo juvenil y aumentar "
    "el valor agregado de las exportaciones.")

speaker(2,
    "Propuesta 2 — Agroindustria y Exportación: otorgar subsidios e infraestructura logística "
    "para diversificar hacia el aguacate Hass y el plátano tecnificado, implementando cadenas "
    "de frío comunitarias. Meta: reducir la dependencia del café y aumentar un 20% las "
    "exportaciones no tradicionales.")

speaker(3,
    "Propuesta 3 — Infraestructura Sostenible: inversión masiva en acueductos rurales veredales "
    "y mejoramiento de vías terciarias mediante placas-huella en los municipios del occidente. "
    "Esto cerraría la brecha urbano-rural y reduciría la pobreza multidimensional en las "
    "comunidades más vulnerables.")

speaker(4,
    "Propuesta 4 — Turismo Sostenible de Alto Valor: certificación en sostenibilidad e incentivos "
    "fiscales para cadenas de ecoturismo y etnoturismo en municipios intermedios. Descentralizaría "
    "la riqueza hacia la periferia del departamento, crearía empleo formal rural y protegería "
    "el Paisaje Cultural Cafetero como activo económico sostenible.")

note("Navegar entre las cuatro pestañas del panel para mostrar cada propuesta.")

# ════════════════════════════════════════
#  SECCIÓN 06 — CONCLUSIONES
# ════════════════════════════════════════
section_tag(6, "Conclusiones", DORADO)
divider(DORADO)

speaker(1,
    "Llegamos a las conclusiones. Risaralda demuestra ser un departamento resiliente con una "
    "ubicación geográfica estratégica óptima para la logística nacional, consolidándose como "
    "un modelo de desarrollo regional intermedio en Colombia.")

speaker(2,
    "Sus fortalezas son claras: alta competitividad logística, indicadores sociales superiores "
    "al promedio nacional — IDH 0,764, pobreza monetaria 22,1%, Gini 0,464 —, liderazgo "
    "exportador no minero-energético y fuerte cohesión gremial e institucional.",
    "IDH 0,764  ·  Pobreza 22,1%  ·  Gini 0,464  ·  Exportaciones +54,6% (2019–2023)")

speaker(3,
    "Las debilidades persisten: alta concentración económica en Pereira-Dosquebradas, "
    "informalidad laboral del 41,5%, brecha urbano-rural en servicios públicos y "
    "vulnerabilidad ante la volatilidad del precio internacional del café.")

speaker(4,
    "El mayor reto es lograr la industrialización tecnológica y la descentralización de "
    "las inversiones para blindar la economía rural ante choques externos, diversificar "
    "la base productiva y garantizar el acceso equitativo a servicios en todo el territorio.")

speaker(1,
    "En síntesis, Risaralda tiene las bases para convertirse en un referente de desarrollo "
    "regional sostenible en Colombia. Las cuatro propuestas apuntan en esa dirección, y la "
    "evidencia de los indicadores muestra que el camino ya está trazado.")

speaker(2, "Muchas gracias por su atención. Quedamos disponibles para preguntas.")

note("Fin de la presentación. Mostrar la sección de conclusiones en pantalla.")

# ════════════════════════════════════════
#  TABLA DE DATOS CLAVE
# ════════════════════════════════════════
doc.add_page_break()
heading("Tabla de datos clave — Referencia rápida", 1, VERDE)
divider()

filas = [
    ("Indicador",                   "Risaralda",    "Colombia",     "Fuente / Año"),
    ("Población",                   "~997.000",     "~52.200.000",  "DANE 2025"),
    ("Densidad",                    "241 hab/km²",  "48 hab/km²",   "DANE 2025"),
    ("Urbanización",                "78%",          "78%",          "CNPV 2018"),
    ("IDH",                         "0,764",        "0,754",        "PNUD 2024"),
    ("Esperanza de vida",           "78,1 años",    "77,5 años",    "DANE 2024"),
    ("Pobreza monetaria",           "22,1%",        "32,7%",        "DANE 2024"),
    ("Pobreza multidimensional",    "8,2%",         "9,5%",         "DANE 2024"),
    ("Coeficiente de Gini",         "0,464",        "0,541",        "DANE 2024"),
    ("Cobertura salud (SGSSS)",     "96,8%",        "~95%",         "SGSSS 2024"),
    ("Analfabetismo",               "3,4%",         "4,7%",         "DANE 2024"),
    ("PIB (% del nacional)",        "1,65%",        "—",            "DANE 2024"),
    ("PIB per cápita",              "USD 6.380",    "USD 6.700",    "DANE 2024"),
    ("Desempleo AM Pereira",        "9,8%",         "~9,2%",        "DANE 2025"),
    ("Informalidad laboral",        "41,5%",        "~55%",         "DANE 2024"),
    ("Exportaciones",               "USD ~342 M",   "—",            "DANE 2023"),
    ("IPC anual 2024",              "5,05%",        "5,20%",        "BanRep 2024"),
    ("IPC anual 2025 (p)",          "4,18%",        "4,32%",        "BanRep 2025"),
    ("Visitantes turismo/año",      "~1,2 M",       "—",            "MinCIT 2023"),
    ("Turismo / PIB dept.",         "~4,5%",        "~2,5%",        "DANE 2022"),
]

table = doc.add_table(rows=len(filas), cols=4)
table.style = "Table Grid"

for i, row_data in enumerate(filas):
    row     = table.rows[i]
    is_head = (i == 0)
    is_even = (i % 2 == 0) and not is_head
    for j, cell in enumerate(row.cells):
        cell.text = ""
        p2 = cell.paragraphs[0]
        r2 = p2.add_run(row_data[j])
        txt_color = BLANCO if is_head else (VERDE if j == 0 else NEGRO)
        font(r2, 9.5 if not is_head else 10, bold=is_head, color=txt_color)
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(2)
        if is_head:
            cell_shade(cell, hex6(VERDE))
        elif is_even:
            cell_shade(cell, "F0EDE8")

doc.add_paragraph()
divider(DORADO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Equipo Asesor: Alejandro Aguilar · Alejandro Cabrera · Luis Pertuz · Jesus Ramirez\n"
    "Fundamentos de Economía · Universidad Libre · 2026\n"
    "Fuentes: DANE · PNUD · Banco de la República · MinCIT · ProColombia · Gobernación de Risaralda"
)
font(r, 8.5, italic=True, color=GRIS)

ruta = r"c:\Users\chuch\Documents\EconomiaRisaralda\Guion_Presentacion_Risaralda.docx"
doc.save(ruta)
print(f"Archivo guardado en:\n  {ruta}")
